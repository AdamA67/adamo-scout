import streamlit as st
from agent import run_pipeline
import json

st.set_page_config(
    page_title="Adamo Scout",
    page_icon="🎯",
    layout="wide"
)

st.title("🎯 Adamo Scout")
st.markdown("*Autonomous candidate research agent — powered by Claude*")
st.divider()

job_brief = st.text_area(
    "Paste the job brief or role description here",
    height=200,
    placeholder="e.g. We are looking for a Director-level Technology Recruitment specialist with 8+ years experience, currently billing £500k+, ideally from a mid-size London tech recruiter..."
)

run_button = st.button("Run Scout", type="primary", use_container_width=True)

if run_button:
    if not job_brief.strip():
        st.error("Please paste a job brief first.")
    else:
        with st.spinner("Analyst is reading the brief..."):
            try:
                # Run the full pipeline with live status updates
                import anthropic
                import os
                from dotenv import load_dotenv
                load_dotenv()

                results = {}

                # Progress bar
                progress = st.progress(0, text="Starting pipeline...")

                progress.progress(10, text="🔍 Analyst reading brief...")
                from agent import load_skill, run_skill

                analyst_output = run_skill(
                    "analyst",
                    f"Here is the job brief:\n\n{job_brief}"
                )
                results["analyst"] = analyst_output
                progress.progress(35, text="📊 Researcher gathering market intel...")

                researcher_output = run_skill(
                    "researcher",
                    "Enrich this candidate profile with market intelligence.",
                    context=f"Analyst output:\n{json.dumps(analyst_output, indent=2)}"
                )
                results["researcher"] = researcher_output
                progress.progress(60, text="♟️ Strategist building search plan...")

                strategist_output = run_skill(
                    "strategist",
                    "Build the search strategy based on this data.",
                    context=f"Analyst output:\n{json.dumps(analyst_output, indent=2)}\n\nResearcher output:\n{json.dumps(researcher_output, indent=2)}"
                )
                results["strategist"] = strategist_output
                progress.progress(85, text="✍️ Writer crafting outreach messages...")

                writer_output = run_skill(
                    "writer",
                    "Write the outreach messages for this candidate profile.",
                    context=f"Analyst output:\n{json.dumps(analyst_output, indent=2)}\n\nResearcher output:\n{json.dumps(researcher_output, indent=2)}\n\nStrategist output:\n{json.dumps(strategist_output, indent=2)}"
                )
                results["writer"] = writer_output
                progress.progress(92, text="🎯 Scorer evaluating search difficulty...")

                scorer_output = run_skill(
                    "scorer",
                    "Evaluate this search and provide the executive summary.",
                    context=f"Analyst output:\n{json.dumps(analyst_output, indent=2)}\n\nResearcher output:\n{json.dumps(researcher_output, indent=2)}\n\nStrategist output:\n{json.dumps(strategist_output, indent=2)}\n\nWriter output:\n{json.dumps(writer_output, indent=2)}"
                )
                results["scorer"] = scorer_output
                
                progress.progress(100, text="✅ Done!")

                st.success("Pipeline complete!")
                st.divider()

                # --- SCORER SUMMARY (shown first) ---
                sc = results["scorer"]

                score = sc.get('search_difficulty_score', 0)
                if score <= 3:
                    score_color = "🟢"
                elif score <= 6:
                    score_color = "🟡"
                else:
                    score_color = "🔴"

                st.subheader("📋 Executive Summary")
                st.info(sc.get('executive_summary', ''))

                metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
                with metric_col1:
                    st.metric("Search Difficulty", f"{score_color} {score}/10")
                with metric_col2:
                    st.metric("Time to Place", sc.get('estimated_time_to_place', ''))
                with metric_col3:
                    st.metric("Confidence", sc.get('confidence_level', ''))
                with metric_col4:
                    st.metric("Biggest Risk", "⚠️ See below")

                st.warning(f"**Biggest Risk:** {sc.get('biggest_risk', '')}")

                st.markdown("**First 3 Moves This Week:**")
                for i, move in enumerate(sc.get('top_3_first_moves', []), 1):
                    st.markdown(f"**{i}.** {move}")

                st.divider()

                # --- RESULTS DISPLAY ---

                col1, col2 = st.columns(2)

                with col1:
                    st.subheader("🔍 Candidate Profile")
                    a = results["analyst"]
                    st.markdown(f"**Role:** {a.get('job_title', '')}")
                    st.markdown(f"**Seniority:** {a.get('seniority_level', '')}")
                    st.markdown(f"**Sector:** {a.get('sector', '')}")
                    st.markdown(f"**Candidate Summary:** {a.get('candidate_summary', '')}")
                    st.markdown("**Key Skills:**")
                    for skill in a.get('key_skills', []):
                        st.markdown(f"- {skill}")
                    st.markdown("**Likely Current Employers:**")
                    for emp in a.get('likely_current_employers', []):
                        st.markdown(f"- {emp}")

                with col2:
                    st.subheader("📊 Market Intelligence")
                    r = results["researcher"]
                    st.markdown(f"**Market Insights:** {r.get('market_insights', '')}")
                    st.markdown(f"**Compensation Range:** {r.get('compensation_range', '')}")
                    st.markdown("**Candidate Move Signals:**")
                    for signal in r.get('candidate_signals', []):
                        st.markdown(f"- {signal}")
                    st.markdown("**Search Keywords:**")
                    st.code(" | ".join(r.get('search_keywords', [])))

                st.divider()

                col3, col4 = st.columns(2)

                with col3:
                    st.subheader("♟️ Search Strategy")
                    s = results["strategist"]
                    st.markdown(f"**Strategy:** {s.get('search_strategy', '')}")
                    st.markdown("**Boolean Search String:**")
                    st.code(s.get('boolean_search_string', ''))
                    st.markdown("**Red Flags to Screen Out:**")
                    for flag in s.get('red_flags', []):
                        st.markdown(f"- {flag}")
                    st.markdown("**Target Companies (Ranked):**")
                    for company in s.get('target_employer_list', []):
                        if isinstance(company, dict):
                            for k, v in company.items():
                                st.markdown(f"**{k}:** {v}")
                        else:
                            st.markdown(f"- {company}")

                with col4:
                    st.subheader("✍️ Outreach Messages")
                    w = results["writer"]
                    st.markdown(f"**Subject Line:** `{w.get('subject_line', '')}`")
                    st.markdown("**Initial Outreach:**")
                    st.info(w.get('outreach_message', ''))
                    st.markdown("**Follow-up (Day 5):**")
                    st.info(w.get('follow_up_message', ''))

                st.divider()

                # Download button
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io

                doc = Document()

                # Title
                title = doc.add_heading('ADAMO SCOUT — CANDIDATE BRIEF', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Scorer section
                doc.add_heading('Executive Summary', 1)
                doc.add_paragraph(results['scorer'].get('executive_summary', ''))
                doc.add_paragraph(f"Search Difficulty: {results['scorer'].get('search_difficulty_score', '')}/10")
                doc.add_paragraph(f"Estimated Time to Place: {results['scorer'].get('estimated_time_to_place', '')}")
                doc.add_paragraph(f"Confidence Level: {results['scorer'].get('confidence_level', '')}")
                doc.add_paragraph(f"Biggest Risk: {results['scorer'].get('biggest_risk', '')}")
                doc.add_heading('First 3 Moves', 2)
                for move in results['scorer'].get('top_3_first_moves', []):
                    doc.add_paragraph(move, style='List Bullet')
                # Analyst section
                doc.add_heading('Candidate Profile', 1)
                doc.add_paragraph(f"Role: {results['analyst'].get('job_title', '')}")
                doc.add_paragraph(f"Seniority: {results['analyst'].get('seniority_level', '')}")
                doc.add_paragraph(f"Sector: {results['analyst'].get('sector', '')}")
                doc.add_paragraph(f"Summary: {results['analyst'].get('candidate_summary', '')}")

                doc.add_heading('Key Skills', 2)
                for skill in results['analyst'].get('key_skills', []):
                    doc.add_paragraph(skill, style='List Bullet')

                doc.add_heading('Likely Current Employers', 2)
                for emp in results['analyst'].get('likely_current_employers', []):
                    doc.add_paragraph(emp, style='List Bullet')

                # Researcher section
                doc.add_heading('Market Intelligence', 1)
                doc.add_paragraph(results['researcher'].get('market_insights', ''))
                doc.add_paragraph(f"Compensation Range: {results['researcher'].get('compensation_range', '')}")

                doc.add_heading('Candidate Move Signals', 2)
                for signal in results['researcher'].get('candidate_signals', []):
                    doc.add_paragraph(signal, style='List Bullet')

                doc.add_heading('Search Keywords', 2)
                doc.add_paragraph(" | ".join(results['researcher'].get('search_keywords', [])))

                # Strategist section
                doc.add_heading('Search Strategy', 1)
                doc.add_paragraph(results['strategist'].get('search_strategy', ''))

                doc.add_heading('Boolean Search String', 2)
                doc.add_paragraph(results['strategist'].get('boolean_search_string', ''))

                doc.add_heading('Red Flags', 2)
                for flag in results['strategist'].get('red_flags', []):
                    doc.add_paragraph(flag, style='List Bullet')

                doc.add_heading('Target Companies', 2)
                for company in results['strategist'].get('target_employer_list', []):
                    if isinstance(company, dict):
                        for k, v in company.items():
                            doc.add_paragraph(f"{k}: {v}", style='List Bullet')
                    else:
                        doc.add_paragraph(company, style='List Bullet')

                # Writer section
                doc.add_heading('Outreach Messages', 1)
                doc.add_heading('Subject Line', 2)
                doc.add_paragraph(results['writer'].get('subject_line', ''))
                doc.add_heading('Initial Message', 2)
                doc.add_paragraph(results['writer'].get('outreach_message', ''))
                doc.add_heading('Follow-up Message (Day 5)', 2)
                doc.add_paragraph(results['writer'].get('follow_up_message', ''))

                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📥 Download Full Brief (.docx)",
                    data=buffer,
                    file_name=f"adamo_scout_{results['analyst'].get('job_title', 'brief').replace(' ', '_').lower()}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            except json.JSONDecodeError as e:
                st.error(f"A skill returned invalid JSON. Try running again. Error: {e}")
            except Exception as e:
                st.error(f"Something went wrong: {e}")